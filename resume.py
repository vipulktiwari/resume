from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "VipulKumarTiwari.docx"

# ── Color palette ─────────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x1B, 0x26, 0x31)
ACCENT      = RGBColor(0x21, 0x8B, 0xC2)    # steel blue
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
NEAR_BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
MUTED       = RGBColor(0x60, 0x60, 0x60)
LIGHT_GRAY  = RGBColor(0x88, 0x88, 0x88)
SIDEBAR_BG  = "F2F6FA"
RULE_COLOR  = "218BC2"

DOC_BG      = "FFFFFF"

# ── Fonts ─────────────────────────────────────────────────────────────────────
FONT_MAIN   = "Calibri"
FONT_HEADER = "Calibri Light"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin   = Inches(0.5)
    section.right_margin  = Inches(0.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def _remove_default_para(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)


def set_run(run, text, bold=False, size=9, color=NEAR_BLACK, font=FONT_MAIN):
    run.text = text
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color


def add_paragraph(container, text, bold=False, size=9, color=NEAR_BLACK,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=2):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_run(run, text, bold=bold, size=size, color=color)
    return p


def _border_bottom(p, color_hex, sz="8"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_section_header(container, text, accent=False):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = FONT_MAIN
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT if accent else DARK_NAVY
    _border_bottom(p, RULE_COLOR if accent else "1B2631", sz="6")
    return p


def add_bullet(container, text, size=8.5):
    import re
    p = container.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.18)
    for part, is_bold in re.findall(r'\*\*(.+?)\*\*|([^*]+)', text):
        segment = part if part else is_bold
        if not segment:
            continue
        run = p.add_run(segment)
        run.bold = bool(part)
        run.font.name = FONT_MAIN
        run.font.size = Pt(size)
        run.font.color.rgb = NEAR_BLACK


def add_job_header(container, title, company, date, location):
    # Title (bold, left)  —  Date (regular, right)
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.RIGHT)

    r_title = p.add_run(title)
    r_title.bold = True
    r_title.font.name = FONT_MAIN
    r_title.font.size = Pt(9.5)
    r_title.font.color.rgb = NEAR_BLACK

    p.add_run("\t")
    r_date = p.add_run(date)
    r_date.font.name = FONT_MAIN
    r_date.font.size = Pt(8.5)
    r_date.font.color.rgb = LIGHT_GRAY
    r_date.italic = True

    # Company (bold, left)  —  Location (regular, right)
    p2 = container.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    p2.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.RIGHT)

    r_comp = p2.add_run(company)
    r_comp.bold = True
    r_comp.font.name = FONT_MAIN
    r_comp.font.size = Pt(8.5)
    r_comp.font.color.rgb = MUTED

    p2.add_run("\t")
    r_loc = p2.add_run(location)
    r_loc.font.name = FONT_MAIN
    r_loc.font.size = Pt(8.5)
    r_loc.font.color.rgb = LIGHT_GRAY
    r_loc.italic = True


def set_cell_borders_none(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_width(cell, inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def set_table_width(table, inches):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(inches * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def set_table_layout_fixed(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_cell_padding(cell, top=0, bottom=0, left=72, right=72):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def add_horizontal_rule(container, color_hex=RULE_COLOR, sz="12", space_before=0, space_after=6):
    """Add a thick horizontal accent line."""
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    # Add zero-width space so line renders
    run = p.add_run("")
    return p


# ════════════════════════════════════════════════════════════
# HEADER
# ════════════════════════════════════════════════════════════
hdr_table = doc.add_table(rows=1, cols=1)
hdr_cell  = hdr_table.cell(0, 0)
shade_cell(hdr_cell, DOC_BG)
set_cell_borders_none(hdr_cell)
set_cell_padding(hdr_cell, top=120, bottom=0, left=0, right=0)

_remove_default_para(hdr_cell)

# Name
p_name = hdr_cell.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_name.paragraph_format.space_before = Pt(2)
p_name.paragraph_format.space_after  = Pt(2)
r = p_name.add_run("Vipul Kumar Tiwari")
r.bold = True
r.font.name = FONT_HEADER
r.font.size = Pt(26)
r.font.color.rgb = DARK_NAVY

# Title
p_title = hdr_cell.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_before = Pt(2)
p_title.paragraph_format.space_after  = Pt(4)
r2 = p_title.add_run("Principal Member of Technical Staff")
r2.font.name = FONT_MAIN
r2.font.size = Pt(12)
r2.font.color.rgb = DARK_NAVY

# Contact line with separators
p_contact = hdr_cell.add_paragraph()
p_contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_contact.paragraph_format.space_before = Pt(2)
p_contact.paragraph_format.space_after  = Pt(2)

contacts = [
    ("+91-8839814859", False),
    ("  •  ", False),
    ("er.vktcs@gmail.com", False),
    ("  •  ", False),
    ("linkedin.com/in/vipulkumartiwari", False),
]
for txt, bold in contacts:
    r3 = p_contact.add_run(txt)
    r3.bold = bold
    r3.font.name = FONT_MAIN
    r3.font.size = Pt(9)
    r3.font.color.rgb = MUTED if "•" in txt else ACCENT

# Thick accent rule under header
#add_horizontal_rule(hdr_cell, color_hex=RULE_COLOR, sz="16", space_before=4, space_after=0)

# ════════════════════════════════════════════════════════════
# SUMMARY
# ════════════════════════════════════════════════════════════
sum_table = doc.add_table(rows=1, cols=1)
sum_cell  = sum_table.cell(0, 0)
set_table_width(sum_table, 7.5)
set_table_layout_fixed(sum_table)
set_cell_width(sum_cell, 7.5)
set_cell_borders_none(sum_cell)
set_cell_padding(sum_cell, top=40, bottom=40, left=0, right=0)

_remove_default_para(sum_cell)

add_section_header(sum_cell, "Professional Summary")
add_paragraph(sum_cell,
    "Software Developer with 10+ years of experience in designing and building scalable, "
    "high-performance products. Expert in system design, problem-solving, and developing "
    "efficient, maintainable code. Extensive experience in databases and storage systems, "
    "with a strong focus on performance optimization, scalability, and reliability of "
    "data-driven applications.",
    size=9, space_before=4, space_after=4, color=NEAR_BLACK)

# ════════════════════════════════════════════════════════════
# BODY  (two-column table: left main, right sidebar)
# ════════════════════════════════════════════════════════════
body_table = doc.add_table(rows=1, cols=2)
left  = body_table.cell(0, 0)
right = body_table.cell(0, 1)

set_table_width(body_table, 7.5)
set_table_layout_fixed(body_table)
set_cell_borders_none(left)
set_cell_borders_none(right)
set_cell_width(left,  4.6)
set_cell_width(right, 2.9)
set_cell_padding(left,  top=0, bottom=0, left=0, right=90)
set_cell_padding(right, top=0, bottom=0, left=90, right=0)
shade_cell(right, SIDEBAR_BG)

_remove_default_para(left)
_remove_default_para(right)

# ────────────────────────────────────────────────────────────
# LEFT: Experience
# ────────────────────────────────────────────────────────────
add_section_header(left, "Experience")

add_job_header(left, "Principal Member of Technical Staff", "Oracle", "02/2026 – Present", "Pune, India")
add_bullet(left, "Responsible for enhancing the LogMiner component in Oracle RDBMS, used in High Availability (HA) solutions and Oracle GoldenGate for database replication and migration.")

add_job_header(left, "Senior Software Developer", "SAP Labs", "07/2020 – 02/2026", "Pune, India")
add_bullet(left, "Contributed to SAP Adaptive Server Enterprise (high-performance OLTP DB), focusing on kernel-layer enhancements and improving the Job Scheduler subsystem for better performance and reliability.")
add_bullet(left, "Handled high-priority, business-critical escalations for enterprise SAP ASE customers, leveraging **strong expertise in database internals**.")
add_bullet(left, "Designed and developed the Hanaservice Sync Agent, a **Python-based microservice on Kubernetes**, enabling reliable cross-cluster synchronization of Hanaservice Custom Resources (CRs).")
add_bullet(left, "Owned end-to-end delivery of cloud-plane services: component testing frameworks, integration/E2E test suites, observability (alerts & monitoring), and customer-facing documentation.")

add_job_header(left, "Staff Software Engineer", "Druva Data Solutions", "03/2018 – 07/2020", "Pune, India")
add_bullet(left, "Designed and built Quaere, a metadata search microservice, from scratch using **AWS DynamoDB and S3**, enabling efficient and reliable metadata search at scale.")
add_bullet(left, "Enhanced mstore, a mail-storage service leveraging Quaere as the underlying storage namespace provider, improving system efficiency and integration.")
add_bullet(left, "Optimized search performance and reduced COGS by improving underlying architecture and resource utilization.")
add_bullet(left, "Designed a CI system in Python/Docker enabling automated testing and seamless code integration for Quaere.")

add_job_header(left, "Software Developer", "Amdocs DVCI", "09/2015 – 03/2018", "Pune, India")
add_bullet(left, "Full lifecycle development of Change Requests for Accounts Receivable and Billing modules for AT&T Telecom.")
add_bullet(left, "Provided on-site production support in Mexico for a Revenue Recognition change request, resolving critical issues.")
add_bullet(left, "Developed solutions in the Ensemble MPS component to process Call Detail Records (CDRs) into billing/charge data.")
add_bullet(left, "Delivered multiple Change Requests for T-Mobile US as an MAF/MPS developer, owning DCUT activities end-to-end.")

# ────────────────────────────────────────────────────────────
# RIGHT: Key Achievements
# ────────────────────────────────────────────────────────────
add_section_header(right, "Key Achievements", accent=True)
add_paragraph(right, "Awards & Recognition", bold=True, size=8.5, space_before=4, space_after=2)
add_bullet(right, "Multiple on-the-spot awards from SAP ASE customers for exceptional fix delivery, responsiveness, and professionalism.")
add_bullet(right, "**Outstanding Achievement Award** at Druva for designing and developing Quaere.")
add_bullet(right, "**All India Rank 2451** with a score of 565 in the GATE (CS).")

# ────────────────────────────────────────────────────────────
# RIGHT: Skills
# ────────────────────────────────────────────────────────────
add_section_header(right, "Skills", accent=True)

skill_groups = [
    ("Core",          "System Design · Databases · Storage"),
    ("Languages",     "C/C++ · Python · Golang"),
    ("Cloud / Ops",   "AWS · Docker · Kubernetes"),
    ("DB / Storage",  "DynamoDB · S3"),
    ("CS Foundations","Data Structures · Algorithms"),
]
for label, skills in skill_groups:
    p = right.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(1)
    rl = p.add_run(label + ":  ")
    rl.bold = True
    rl.font.name = FONT_MAIN
    rl.font.size = Pt(8.5)
    rl.font.color.rgb = DARK_NAVY
    rs = p.add_run(skills)
    rs.font.name = FONT_MAIN
    rs.font.size = Pt(8.5)
    rs.font.color.rgb = NEAR_BLACK

# ────────────────────────────────────────────────────────────
# RIGHT: Education
# ────────────────────────────────────────────────────────────
add_section_header(right, "Education", accent=True)

add_paragraph(right, "B.E. Computer Science & Engineering", bold=True, size=9, space_before=5, space_after=1)
add_paragraph(right, "Rajiv Gandhi Technical University", size=8, color=MUTED, space_before=0, space_after=1)
add_paragraph(right, "2011 – 2015  ·  Bhopal", size=8, color=MUTED, space_before=0, space_after=1)
add_paragraph(right, "CGPA: 8.07 / 10", size=8.5, bold=True, space_before=0, space_after=6)

add_paragraph(right, "AISSCE (Class XII)", bold=True, size=9, space_before=4, space_after=1)
add_paragraph(right, "Jawahar Navodaya Vidhyalaya, Narsinghpur", size=8, color=MUTED, space_before=0, space_after=1)
add_paragraph(right, "2009 – 2010", size=8, color=MUTED, space_before=0, space_after=1)
add_paragraph(right, "Score: 79.6 / 100", size=8.5, bold=True, space_before=0, space_after=6)

add_paragraph(right, "AISSE (Class X)", bold=True, size=9, space_before=4, space_after=1)
add_paragraph(right, "Jawahar Navodaya Vidhyalaya, Narsinghpur", size=8, color=MUTED, space_before=0, space_after=1)
add_paragraph(right, "2007-2008", size=8, color=MUTED, space_before=0, space_after=1)
add_paragraph(right, "Score: 81.6 / 100", size=8.5, bold=True, space_before=0, space_after=6)

# ────────────────────────────────────────────────────────────
# RIGHT: Languages
# ────────────────────────────────────────────────────────────
add_section_header(right, "Languages", accent=True)
add_paragraph(right, "Hindi", bold=True, size=8.5, space_before=4, space_after=1)
add_paragraph(right, "Native proficiency", size=8, color=MUTED, space_before=0, space_after=4)
add_paragraph(right, "English", bold=True, size=8.5, space_before=0, space_after=1)
add_paragraph(right, "Professional proficiency", size=8, color=MUTED, space_before=0, space_after=4)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
